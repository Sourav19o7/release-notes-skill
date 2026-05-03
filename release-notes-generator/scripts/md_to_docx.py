#!/usr/bin/env python3
"""
Convert Markdown release notes to Word document (.docx)

Usage:
    python md_to_docx.py input.md output.docx
    python md_to_docx.py input.md  # outputs to input.docx
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table and return rows."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = lines[i].strip()
        # Skip separator row (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', row):
            i += 1
            continue
        # Parse cells
        cells = [cell.strip() for cell in row.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # Bold header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def extract_link(text):
    """Extract markdown links and return plain text."""
    # Convert [text](url) to text
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Handle headings
        if stripped.startswith('# '):
            heading = doc.add_heading(stripped[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)

        # Handle tables
        elif stripped.startswith('|'):
            rows, i = parse_markdown_table(lines, i)
            add_table_to_doc(doc, rows)
            continue  # i is already incremented

        # Handle blockquotes
        elif stripped.startswith('> '):
            text = extract_link(stripped[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Handle bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = extract_link(stripped[2:])
            doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = extract_link(re.sub(r'^\d+\.\s', '', stripped))
            doc.add_paragraph(text, style='List Number')

        # Handle bold text markers (like **text**)
        elif stripped.startswith('**') and stripped.endswith('**'):
            text = stripped[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True

        # Handle horizontal rules
        elif stripped in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)

        # Handle code blocks
        elif stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

        # Regular paragraph
        else:
            text = extract_link(stripped)
            doc.add_paragraph(text)

        i += 1

    # Save document
    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])

    if not md_path.exists():
        print(f"Error: {md_path} not found")
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix('.docx')

    convert_md_to_docx(md_path, docx_path)


if __name__ == '__main__':
    main()
